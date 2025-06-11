"""
Stage 4: Text Extraction Service
Extract and process text from various document formats for creative RAG
Part of knowNothing Creative RAG
"""

import os
import re
import logging
from pathlib import Path
from typing import Dict, List, Optional, Any, Tuple
from datetime import datetime
import sqlite3

# Import text extraction libraries
try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

logger = logging.getLogger(__name__)

class TextExtractor:
    """
    Extracts and processes text from various document formats
    Designed for creative professionals working with scripts, notes, and research
    """
    
    def __init__(self, db_path: str = "data/documents.db"):
        self.db_path = db_path
        self._init_text_storage()
        
        # Log available extractors
        self._log_available_extractors()
    
    def _init_text_storage(self):
        """Initialize text storage table in SQLite"""
        try:
            conn = sqlite3.connect(self.db_path)
            cursor = conn.cursor()
            
            cursor.execute("""
                CREATE TABLE IF NOT EXISTS document_text (
                    document_id TEXT PRIMARY KEY,
                    extracted_text TEXT NOT NULL,
                    extraction_method TEXT NOT NULL,
                    word_count INTEGER,
                    character_count INTEGER,
                    page_count INTEGER DEFAULT NULL,
                    extraction_date TEXT NOT NULL,
                    processing_notes TEXT,
                    text_metadata JSONB,
                    FOREIGN KEY (document_id) REFERENCES documents (id)
                )
            """)
            
            # Add extraction status to documents table if not exists
            cursor.execute("""
                ALTER TABLE documents 
                ADD COLUMN text_extracted BOOLEAN DEFAULT FALSE
            """)
            
            conn.commit()
            conn.close()
            logger.info("✅ Text extraction database initialized")
            
        except sqlite3.OperationalError as e:
            if "duplicate column name" in str(e):
                # Column already exists, that's fine
                pass
            else:
                logger.error(f"❌ Text storage initialization failed: {e}")
                raise
        except Exception as e:
            logger.error(f"❌ Text storage initialization failed: {e}")
            raise
    
    def _log_available_extractors(self):
        """Log which text extractors are available"""
        extractors = []
        if PDF_AVAILABLE:
            extractors.append("PDF (PyPDF2)")
        if DOCX_AVAILABLE:
            extractors.append("DOCX (python-docx)")
        extractors.append("TXT (built-in)")
        
        logger.info(f"📄 Available text extractors: {', '.join(extractors)}")
    
    async def extract_text_from_document(self, document_id: str) -> Dict[str, Any]:
        """
        Extract text from a document by ID
        
        Args:
            document_id: Document ID from storage
            
        Returns:
            Dict containing extracted text and metadata
        """
        try:
            # Get document info
            doc_info = await self._get_document_info(document_id)
            if not doc_info:
                raise ValueError(f"Document {document_id} not found")
            
            file_path = Path(doc_info['file_path'])
            if not file_path.exists():
                raise ValueError(f"Document file not found: {file_path}")
            
            file_type = doc_info['file_type'].lower()
            
            # Extract text based on file type
            if file_type == '.pdf':
                result = await self._extract_pdf_text(file_path)
            elif file_type in ['.docx', '.doc']:
                result = await self._extract_docx_text(file_path)
            elif file_type in ['.txt', '.rtf']:
                result = await self._extract_txt_text(file_path)
            else:
                raise ValueError(f"Unsupported file type: {file_type}")
            
            # Store extracted text
            await self._store_extracted_text(document_id, result)
            
            # Update document status
            await self._mark_text_extracted(document_id)
            
            logger.info(f"✅ Text extracted from {doc_info['original_filename']}: {result['word_count']} words")
            
            return {
                "success": True,
                "document_id": document_id,
                "extraction_method": result["method"],
                "text_preview": result["text"][:500] + "..." if len(result["text"]) > 500 else result["text"],
                "statistics": {
                    "word_count": result["word_count"],
                    "character_count": result["character_count"],
                    "page_count": result.get("page_count"),
                    "extraction_quality": result.get("quality", "good")
                },
                "processing_notes": result.get("notes", []),
                "timestamp": datetime.utcnow().isoformat()
            }
            
        except Exception as e:
            logger.error(f"❌ Text extraction failed for {document_id}: {str(e)}")
            return {
                "success": False,
                "error": str(e),
                "document_id": document_id,
                "timestamp": datetime.utcnow().isoformat()
            }
    
    async def _extract_pdf_text(self, file_path: Path) -> Dict[str, Any]:
        """Extract text from PDF file"""
        if not PDF_AVAILABLE:
            raise ValueError("PDF extraction not available. Install PyPDF2: pip install PyPDF2")
        
        try:
            extracted_text = ""
            page_count = 0
            notes = []
            
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                page_count = len(pdf_reader.pages)
                
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text.strip():
                            extracted_text += f"\n--- Page {page_num + 1} ---\n"
                            extracted_text += page_text
                        else:
                            notes.append(f"Page {page_num + 1} appears to be empty or image-only")
                    except Exception as e:
                        notes.append(f"Error extracting page {page_num + 1}: {str(e)}")
            
            # Clean up text
            cleaned_text = self._clean_extracted_text(extracted_text)
            
            return {
                "text": cleaned_text,
                "method": "PyPDF2",
                "page_count": page_count,
                "word_count": len(cleaned_text.split()),
                "character_count": len(cleaned_text),
                "notes": notes,
                "quality": "good" if len(cleaned_text) > 100 else "low"
            }
            
        except Exception as e:
            logger.error(f"❌ PDF extraction failed: {str(e)}")
            raise ValueError(f"Failed to extract PDF text: {str(e)}")
    
    async def _extract_docx_text(self, file_path: Path) -> Dict[str, Any]:
        """Extract text from DOCX file"""
        if not DOCX_AVAILABLE:
            raise ValueError("DOCX extraction not available. Install python-docx: pip install python-docx")
        
        try:
            doc = DocxDocument(str(file_path))
            extracted_text = ""
            notes = []
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    extracted_text += paragraph.text + "\n"
            
            # Extract table content if any
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text.strip() for cell in row.cells])
                    if row_text.strip():
                        extracted_text += row_text + "\n"
                        
            if doc.tables:
                notes.append(f"Extracted content from {len(doc.tables)} table(s)")
            
            # Clean up text
            cleaned_text = self._clean_extracted_text(extracted_text)
            
            return {
                "text": cleaned_text,
                "method": "python-docx",
                "word_count": len(cleaned_text.split()),
                "character_count": len(cleaned_text),
                "notes": notes,
                "quality": "excellent"
            }
            
        except Exception as e:
            logger.error(f"❌ DOCX extraction failed: {str(e)}")
            raise ValueError(f"Failed to extract DOCX text: {str(e)}")
    
    async def _extract_txt_text(self, file_path: Path) -> Dict[str, Any]:
        """Extract text from plain text file"""
        try:
            # Try different encodings
            encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
            extracted_text = ""
            encoding_used = "utf-8"
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        extracted_text = file.read()
                        encoding_used = encoding
                        break
                except UnicodeDecodeError:
                    continue
            
            if not extracted_text:
                raise ValueError("Could not decode text file with any common encoding")
            
            # Clean up text
            cleaned_text = self._clean_extracted_text(extracted_text)
            
            return {
                "text": cleaned_text,
                "method": f"text-file ({encoding_used})",
                "word_count": len(cleaned_text.split()),
                "character_count": len(cleaned_text),
                "notes": [f"Decoded using {encoding_used} encoding"],
                "quality": "excellent"
            }
            
        except Exception as e:
            logger.error(f"❌ Text file extraction failed: {str(e)}")
            raise ValueError(f"Failed to extract text: {str(e)}")
    
    def _clean_extracted_text(self, text: str) -> str:
        """Clean and normalize extracted text"""
        if not text:
            return ""
        
        # Remove excessive whitespace
        text = re.sub(r'\n\s*\n\s*\n', '\n\n', text)  # Multiple newlines to double
        text = re.sub(r'\t+', ' ', text)  # Tabs to spaces
        text = re.sub(r' +', ' ', text)  # Multiple spaces to single
        
        # Remove weird characters but keep creative content
        text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x84\x86-\x9f]', '', text)
        
        # Strip leading/trailing whitespace
        text = text.strip()
        
        return text
    
    async def get_extracted_text(self, document_id: str) -> Optional[Dict[str, Any]]:
        """Get previously extracted text for a document"""
        try:
            conn = sqlite3.connect(self.db_path)
            conn.row_factory = sqlite3.Row
            cursor = conn.cursor()
            
            cursor.execute("""
                SELECT * FROM document_text WHERE document_id = ?
            """, (document_id,))
            
            row = cursor.fetchone()
            conn.close()
            
            if row:
                return {
                    "document_id": row["document_id"],
                    "text": row["extracted_text"],
                    "method": row["extraction_method"],
                    "word_count": row["word_count"],
                    "character_count": row["character_count"],
                    "page_count": row["page_count"],
                    "extraction_date": row["extraction_date"],
                    "processing_notes": row["processing_notes"],
                    "text_metadata": row["text_metadata"]
                }
            
            return None
            
        except Exception as e:
            logger.error(f"❌ Failed to get extracted text: {str(e)}")
            return None
    
    async def get_text_extraction_status(self) -> Dict[str, Any]:
        """Get status of text extraction for all documents"""
        try:
            conn = sqlite3.connect(self.db_path)
            conn.row_factory = sqlite3.Row
            cursor = conn.cursor()
            
            # Get total documents
            cursor.execute("SELECT COUNT(*) as total FROM documents")
            total_docs = cursor.fetchone()["total"]
            
            # Get extracted documents
            cursor.execute("SELECT COUNT(*) as extracted FROM documents WHERE text_extracted = TRUE")
            extracted_docs = cursor.fetchone()["extracted"]
            
            # Get extraction statistics
            cursor.execute("""
                SELECT 
                    extraction_method,
                    COUNT(*) as count,
                    AVG(word_count) as avg_words,
                    SUM(word_count) as total_words
                FROM document_text 
                GROUP BY extraction_method
            """)
            
            method_stats = {}
            for row in cursor.fetchall():
                method_stats[row["extraction_method"]] = {
                    "count": row["count"],
                    "avg_words": round(row["avg_words"] or 0, 1),
                    "total_words": row["total_words"] or 0
                }
            
            conn.close()
            
            return {
                "total_documents": total_docs,
                "extracted_documents": extracted_docs,
                "extraction_percentage": round((extracted_docs / total_docs * 100) if total_docs > 0 else 0, 1),
                "method_statistics": method_stats,
                "available_extractors": {
                    "pdf": PDF_AVAILABLE,
                    "docx": DOCX_AVAILABLE,
                    "txt": True
                }
            }
            
        except Exception as e:
            logger.error(f"❌ Failed to get extraction status: {str(e)}")
            return {"error": str(e)}
    
    async def _get_document_info(self, document_id: str) -> Optional[Dict]:
        """Get document information from storage"""
        try:
            conn = sqlite3.connect(self.db_path)
            conn.row_factory = sqlite3.Row
            cursor = conn.cursor()
            
            cursor.execute("SELECT * FROM documents WHERE id = ?", (document_id,))
            row = cursor.fetchone()
            conn.close()
            
            return dict(row) if row else None
            
        except Exception as e:
            logger.error(f"❌ Failed to get document info: {str(e)}")
            return None
    
    async def _store_extracted_text(self, document_id: str, extraction_result: Dict[str, Any]):
        """Store extracted text in database"""
        try:
            conn = sqlite3.connect(self.db_path)
            cursor = conn.cursor()
            
            cursor.execute("""
                INSERT OR REPLACE INTO document_text
                (document_id, extracted_text, extraction_method, word_count, 
                 character_count, page_count, extraction_date, processing_notes, text_metadata)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
            """, (
                document_id,
                extraction_result["text"],
                extraction_result["method"],
                extraction_result["word_count"],
                extraction_result["character_count"],
                extraction_result.get("page_count"),
                datetime.utcnow().isoformat(),
                str(extraction_result.get("notes", [])),
                "{}"  # Placeholder for future metadata
            ))
            
            conn.commit()
            conn.close()
            
        except Exception as e:
            logger.error(f"❌ Failed to store extracted text: {str(e)}")
            raise
    
    async def _mark_text_extracted(self, document_id: str):
        """Mark document as having text extracted"""
        try:
            conn = sqlite3.connect(self.db_path)
            cursor = conn.cursor()
            
            cursor.execute("""
                UPDATE documents 
                SET text_extracted = TRUE 
                WHERE id = ?
            """, (document_id,))
            
            conn.commit()
            conn.close()
            
        except Exception as e:
            logger.error(f"❌ Failed to mark text extracted: {str(e)}")
            raise